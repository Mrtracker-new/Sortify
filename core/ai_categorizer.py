import os
import re
import pickle
import logging
import hashlib
import numpy as np
from pathlib import Path
from collections import OrderedDict
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
from sklearn.feature_extraction.text import TfidfVectorizer  # Changed from CountVectorizer
from sklearn.naive_bayes import MultinomialNB
from sklearn.ensemble import RandomForestClassifier
from sklearn.pipeline import Pipeline
from sklearn.model_selection import train_test_split, cross_val_score
from sklearn.metrics import accuracy_score, classification_report
from datetime import datetime

# Try to import text extraction libraries, but make them optional
try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    logging.warning("PyPDF2 module not found. PDF text extraction will be limited.")
    PYPDF2_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    logging.warning("python-docx module not found. Word document extraction will be limited.")
    PYTHON_DOCX_AVAILABLE = False



# Try to import sentence-transformers, but make it optional
try:
    from sentence_transformers import SentenceTransformer
    from sklearn.metrics.pairwise import cosine_similarity
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    logging.warning("sentence-transformers module not found. Semantic classification will fall back to TF-IDF.")
    SENTENCE_TRANSFORMERS_AVAILABLE = False
except OSError as e:
    # Catch DLL loading errors from PyTorch dependencies
    logging.warning(f"sentence-transformers dependencies failed to load (DLL error): {e}")
    logging.warning("This is often caused by missing Visual C++ Redistributables.")
    logging.warning("Semantic classification will fall back to TF-IDF.")
    SENTENCE_TRANSFORMERS_AVAILABLE = False
except Exception as e:
    logging.warning(f"Unexpected error importing sentence-transformers: {e}")
    logging.warning("Semantic classification will fall back to TF-IDF.")
    SENTENCE_TRANSFORMERS_AVAILABLE = False


# ---------------------------------------------------------------------------
# Built-in bootstrap training corpus
# ---------------------------------------------------------------------------
# One synthetic feature-string per known Sortify category.  These keyword bags
# are intentionally simple — the goal is to give every category a meaningful
# prototype so the TF-IDF / Naive Bayes model can work out of the box without
# any real files.  When real sorted files are available, auto_train_if_needed()
# will progressively upgrade these prototypes.
# ---------------------------------------------------------------------------
_BOOTSTRAP_CORPUS: list[tuple[str, str]] = [
    # documents
    ("invoice receipt payment bill statement tax pdf report",       "documents/pdf"),
    ("word document report essay letter memo rtf docx odt",        "documents/word"),
    ("spreadsheet budget table data rows columns xls xlsx ods csv","documents/excel"),
    ("slide deck presentation pptx ppt odp show keynote",         "documents/powerpoint"),
    ("plain text readme notes log txt md rst",                     "documents/text"),
    # images
    ("photo picture jpg jpeg jfif portrait landscape snapshot",    "images/jpg"),
    ("screenshot png graphic icon ui transparent clip",            "images/png"),
    ("animation gif meme loop frame sprite",                       "images/gif"),
    ("bitmap bmp raster pixel",                                    "images/bmp"),
    ("webp compressed image modern web",                           "images/webp"),
    ("heic heif apple iphone photo",                               "images/heic"),
    ("tiff tif scanned document high quality raw",                 "images/tiff"),
    ("svg vector ai eps illustrator graphic design",               "images/vector"),
    ("raw cr2 nef arw dng camera unprocessed",                     "images/raw"),
    # ai images
    ("chatgpt dalle gpt openai generated ai image",               "ai_images/chatgpt"),
    ("midjourney mj ai art dream generated",                      "ai_images/midjourney"),
    ("stable diffusion sd ai generative img2img",                  "ai_images/stable_diffusion"),
    ("bing ai image creator microsoft generated",                  "ai_images/bing"),
    # videos
    ("video movie film mp4 mkv avi mov hdvideo",                  "videos/movies"),
    ("shorts clip webm reel youtube brief",                        "videos/shorts"),
    ("whatsapp wa chat video mp4 3gp mobile",                      "videos/whatsapp"),
    ("telegram tg channel video clip",                             "videos/telegram"),
    ("instagram ig reels story video",                             "videos/instagram"),
    ("facebook fb watch video mp4",                                "videos/facebook"),
    ("youtube yt download video mkv mp4",                          "videos/youtube"),
    # audio
    ("music mp3 flac wav aac album track song artist",            "audio/music"),
    ("voice recording ogg wma speech memo podcast",               "audio/voice"),
    # code
    ("python script code py function class import module",        "code/python"),
    ("javascript html css web page browser jsx tsx node",         "code/web"),
    ("sql database json xml yaml data schema config",             "code/data"),
    ("shell bash powershell batch ps1 bat sh command",            "code/scripts"),
    # archives
    ("zip archive compressed rar tar 7z gz bundle",               "archives/compressed"),
    ("iso disk image img virtual cd dvd",                         "archives/disk_images"),
    # office
    ("template dotx potx xltx office style",                     "office/templates"),
    ("outlook email pst ost msg mailbox",                          "office/outlook"),
    ("access database accdb mdb query table",                      "office/database"),
]


def _run_with_timeout(func, args=(), kwargs=None, timeout_seconds=10):
    """Execute a function with a timeout limit.
    
    Args:
        func: Function to execute
        args: Positional arguments for the function
        kwargs: Keyword arguments for the function
        timeout_seconds: Maximum time allowed for execution
        
    Returns:
        The function's return value if successful
        
    Raises:
        TimeoutError: If function execution exceeds timeout_seconds
    """
    if kwargs is None:
        kwargs = {}
    
    with ThreadPoolExecutor(max_workers=1) as executor:
        future = executor.submit(func, *args, **kwargs)
        try:
            return future.result(timeout=timeout_seconds)
        except FuturesTimeoutError:
            raise TimeoutError(f"Function {func.__name__} exceeded timeout of {timeout_seconds} seconds")


def compute_file_hash(file_path):
    """Compute a content-based MD5 hash for a file.

    Uses optimised partial hashing for large files to balance speed and
    uniqueness: files under 1 MB are fully hashed; larger files use the
    first 64 KB + last 64 KB + file size.

    This is the **single authoritative implementation** used by both
    ``ContentCache`` and ``AIFileClassifier`` so that the two caches always
    produce identical keys for the same file.

    Args:
        file_path: Path-like object for the file.

    Returns:
        str: MD5 hex-digest of the file content, or ``None`` on error.
    """
    try:
        stat = file_path.stat()
        file_size = stat.st_size

        # Small files (<1 MB) – hash the entire content
        if file_size < 1024 * 1024:
            with open(file_path, 'rb') as f:
                return hashlib.md5(f.read()).hexdigest()

        # Large files – hash first 64 KB + last 64 KB + size
        hash_obj = hashlib.md5()
        with open(file_path, 'rb') as f:
            hash_obj.update(f.read(65536))
            f.seek(max(0, file_size - 65536))
            hash_obj.update(f.read(65536))
            hash_obj.update(str(file_size).encode())

        return hash_obj.hexdigest()
    except Exception as e:
        logging.debug(f"Error computing file hash for {file_path}: {e}")
        return None


class LRUCache:
    """LRU (Least Recently Used) Cache implementation with automatic eviction.
    
    This cache maintains a maximum size and automatically evicts the oldest
    item when the cache exceeds max_size. Uses OrderedDict for efficient
    O(1) operations and proper ordering.
    """
    def __init__(self, max_size=1000):
        """Initialize LRU cache with maximum size.
        
        Args:
            max_size: Maximum number of items to store in cache (default: 1000)
        """
        self.cache = OrderedDict()
        self.max_size = max_size
        self.hits = 0
        self.misses = 0
    
    def get(self, key):
        """Get an item from cache and mark it as recently used.
        
        Args:
            key: Cache key to retrieve
            
        Returns:
            Cached value if found, None otherwise
        """
        if key in self.cache:
            # Move to end to mark as recently used
            self.cache.move_to_end(key)
            self.hits += 1
            return self.cache[key]
        self.misses += 1
        return None
    
    def put(self, key, value):
        """Add or update an item in cache.
        
        If the item exists, it's moved to the end (most recent).
        If cache exceeds max_size, the oldest item is evicted.
        
        Args:
            key: Cache key
            value: Value to store
        """
        if key in self.cache:
            # Update existing item and move to end
            self.cache.move_to_end(key)
        self.cache[key] = value
        
        # Evict oldest item if cache size exceeded
        if len(self.cache) > self.max_size:
            self.cache.popitem(last=False)  # Remove oldest (first) item
    
    def clear(self):
        """Clear all items from cache."""
        self.cache.clear()
        self.hits = 0
        self.misses = 0
    
    def get_stats(self):
        """Get cache statistics.
        
        Returns:
            dict: Dictionary with cache hits, misses, size, and hit rate
        """
        total = self.hits + self.misses
        hit_rate = (self.hits / total * 100) if total > 0 else 0
        return {
            'hits': self.hits,
            'misses': self.misses,
            'size': len(self.cache),
            'max_size': self.max_size,
            'hit_rate': round(hit_rate, 2)
        }


class ContentCache:
    """Cache for extracted content to avoid repeated slow processing.
    
    Uses file content hash as cache key (optimized for large files with partial hashing).
    This prevents repeated expensive extraction operations (PyPDF2, python-docx, etc.) 
    and eliminates race conditions from mtime-based caching.
    """
    def __init__(self, max_size=500):
        """Initialize content cache with LRU eviction.
        
        Args:
            max_size: Maximum number of content items to cache (default: 500)
        """
        self.cache = LRUCache(max_size=max_size)
    
    def get_file_hash(self, file_path):
        """Generate a unique hash for a file based on its content.

        Delegates to the module-level :func:`compute_file_hash` so that
        ``ContentCache`` and ``AIFileClassifier`` always share the same
        hashing strategy and produce identical cache keys.

        Args:
            file_path: Path object for the file.

        Returns:
            str: MD5 hash of file content, or ``None`` on error.
        """
        return compute_file_hash(file_path)
    
    def get_content(self, file_path):
        """Retrieve cached content for a file.
        
        Args:
            file_path: Path object for the file
            
        Returns:
            str: Cached content if available, None otherwise
        """
        file_hash = self.get_file_hash(file_path)
        if file_hash is None:
            return None
        return self.cache.get(file_hash)
    
    def set_content(self, file_path, content):
        """Cache extracted content for a file.
        
        Args:
            file_path: Path object for the file
            content: Extracted content to cache
        """
        file_hash = self.get_file_hash(file_path)
        if file_hash is not None:
            self.cache.put(file_hash, content)
    
    def get_stats(self):
        """Get cache statistics.
        
        Returns:
            dict: Cache statistics including hits, misses, and hit rate
        """
        return self.cache.get_stats()
    
    def clear(self):
        """Clear all cached content."""
        self.cache.clear()


class SentenceTransformerClassifier:
    """Semantic file classifier using Sentence Transformers for embeddings.
    
    This classifier uses the all-MiniLM-L6-v2 model (80MB) to generate semantic
    embeddings of file content, then uses cosine similarity to match files to
    categories. This provides significantly better accuracy than TF-IDF for
    understanding context and meaning in file content.
    """
    
    def __init__(self, model_name='all-MiniLM-L6-v2'):
        """Initialize the Sentence Transformer classifier.
        
        Args:
            model_name: Name of the sentence-transformers model to use
                       (default: 'all-MiniLM-L6-v2' - 80MB, fast, accurate)
        """
        if not SENTENCE_TRANSFORMERS_AVAILABLE:
            raise ImportError(
                "sentence-transformers is not installed. "
                "Install it with: pip install sentence-transformers"
            )
        
        self.model_name = model_name
        self.model = None
        self.category_embeddings = {}  # Category prototype embeddings
        self.categories = []
        self.trained = False
        
        # Cache for embeddings to avoid recomputing
        self.embedding_cache = LRUCache(max_size=1000)
        
        # Initialize the model on first use (lazy loading)
        self._initialize_model()
    
    def _initialize_model(self):
        """Initialize the sentence transformer model (lazy loading)."""
        if self.model is None:
            try:
                logging.info(f"Loading Sentence Transformer model: {self.model_name}")
                logging.info("First time download may take a minute (model is ~80MB)...")
                self.model = SentenceTransformer(self.model_name)
                logging.info(f"Model {self.model_name} loaded successfully!")
            except Exception as e:
                logging.error(f"Failed to load Sentence Transformer model: {e}")
                raise
    
    def _get_embedding(self, text):
        """Generate embedding for text with caching.
        
        Args:
            text: Text to generate embedding for
            
        Returns:
            numpy array: Embedding vector for the text
        """
        # Create cache key using hash of text
        cache_key = hashlib.md5(text.encode('utf-8', errors='ignore')).hexdigest()
        
        # Check cache first
        cached_embedding = self.embedding_cache.get(cache_key)
        if cached_embedding is not None:
            return cached_embedding
        
        # Generate embedding
        embedding = self.model.encode(text, convert_to_numpy=True)
        
        # Cache it
        self.embedding_cache.put(cache_key, embedding)
        
        return embedding
    
    def train(self, texts, categories):
        """Train the classifier on text examples.
        
        Args:
            texts: List of text features (from extract_features)
            categories: List of corresponding categories
        """
        if len(texts) != len(categories):
            raise ValueError("Number of texts and categories must match")
        
        # Store unique categories
        self.categories = list(set(categories))
        
        logging.info(f"Training Sentence Transformer classifier on {len(texts)} samples...")
        
        # For each category, compute average embedding (prototype)
        category_texts = {cat: [] for cat in self.categories}
        for text, category in zip(texts, categories):
            category_texts[category].append(text)
        
        # Compute prototype embeddings
        for category in self.categories:
            texts_for_category = category_texts[category]
            embeddings = [self._get_embedding(text) for text in texts_for_category]
            
            # Average all embeddings for this category
            avg_embedding = np.mean(embeddings, axis=0)
            self.category_embeddings[category] = avg_embedding
            
            logging.debug(f"  Category '{category}': {len(texts_for_category)} samples")
        
        self.trained = True
        logging.info(f"Training complete! {len(self.categories)} categories learned.")
    
    def predict(self, text):
        """Predict category for text with confidence scores.
        
        Args:
            text: Text features to classify
            
        Returns:
            tuple: (predicted_category, confidence_score, probabilities_dict)
        """
        if not self.trained:
            raise ValueError("Classifier must be trained before prediction")
        
        # Generate embedding for the input text
        text_embedding = self._get_embedding(text).reshape(1, -1)
        
        # Calculate cosine similarity with each category prototype
        similarities = {}
        for category, category_embedding in self.category_embeddings.items():
            similarity = cosine_similarity(
                text_embedding, 
                category_embedding.reshape(1, -1)
            )[0][0]
            similarities[category] = float(similarity)
        
        # Get the category with highest similarity
        predicted_category = max(similarities, key=similarities.get)
        confidence = similarities[predicted_category]
        
        # Convert similarities to pseudo-probabilities (softmax-like normalization)
        # This makes them comparable to sklearn's predict_proba output
        total = sum(similarities.values())
        if total > 0:
            probabilities = {cat: sim / total for cat, sim in similarities.items()}
        else:
            probabilities = {cat: 1.0 / len(similarities) for cat in similarities}
        
        return predicted_category, confidence, probabilities
    
    def get_embedding_stats(self):
        """Get statistics about embedding cache.
        
        Returns:
            dict: Cache statistics
        """
        return self.embedding_cache.get_stats()



class AIFileClassifier:
    """AI-based file classifier using Naive Bayes with content analysis"""
    def __init__(self, model_path=None, classifier_type='naive_bayes', max_file_size_mb=100, extraction_timeout_seconds=10):
        # Create a pipeline with improved vectorizer and classifier
        self.classifier_type = classifier_type
        self._create_pipeline()
        self.trained = False
        self.categories = []
        # Use LRU cache with automatic eviction to prevent memory leaks
        self.feature_cache = LRUCache(max_size=1000)
        # Add content cache specifically for extracted content
        self.content_cache = ContentCache(max_size=500)
        self.last_training_time = None
        self.training_accuracy = None
        self.model_version = 1
        
        # Configuration for content extraction processing safety
        self.max_file_size_mb = max_file_size_mb
        self.extraction_timeout_seconds = extraction_timeout_seconds
        
        # Track extraction failures to warn users about degraded categorization
        self.extraction_failures = []  # List of files where content extraction failed
        self.extraction_warnings = {}  # Dict mapping file paths to warning messages
        
        # Bootstrap on the built-in keyword corpus so the model is usable right away.
        # This runs before load_model() so a persisted model (progressive learning)
        # always wins — it simply overwrites the bootstrap state.
        self._bootstrap_train()
        
        # Override with persisted model if available (preserves progressive learning)
        if model_path and os.path.exists(model_path):
            self.load_model(model_path)
            
    def _create_pipeline(self):
        """Create the ML pipeline based on the selected classifier type"""
        if self.classifier_type == 'sentence_transformer':
            # Use Sentence Transformers for semantic understanding
            if SENTENCE_TRANSFORMERS_AVAILABLE:
                self.model = SentenceTransformerClassifier()
                logging.info("Using Sentence Transformer classifier for semantic categorization")
            else:
                logging.warning("Sentence Transformers not available, falling back to Naive Bayes")
                self.classifier_type = 'naive_bayes'
                self.model = Pipeline([
                    ('vectorizer', TfidfVectorizer(analyzer='word', ngram_range=(1, 3), max_features=5000)),
                    ('classifier', MultinomialNB())
                ])
        elif self.classifier_type == 'random_forest':
            self.model = Pipeline([
                ('vectorizer', TfidfVectorizer(analyzer='word', ngram_range=(1, 3), max_features=5000)),
                ('classifier', RandomForestClassifier(n_estimators=100, random_state=42))
            ])
        else:  # Default to naive_bayes
            self.model = Pipeline([
                ('vectorizer', TfidfVectorizer(analyzer='word', ngram_range=(1, 3), max_features=5000)),
                ('classifier', MultinomialNB())
            ])

    
    def _compute_file_hash(self, file_path):
        """Compute content hash for a file.

        Delegates to the module-level :func:`compute_file_hash` so that
        ``AIFileClassifier`` and ``ContentCache`` always produce identical
        cache keys for the same file.

        Args:
            file_path: Path object for the file.

        Returns:
            str: MD5 hash of file content, or ``None`` on error.
        """
        return compute_file_hash(file_path)
    
    def _extract_pdf_text(self, file_path):
        """Extract text from PDF file using PyPDF2.
        
        Args:
            file_path: Path object for the PDF file
            
        Returns:
            str: Extracted text content, or None on error
        """
        if not PYPDF2_AVAILABLE:
            return None
            
        try:
            text_parts = []
            with open(file_path, 'rb') as f:
                pdf_reader = PdfReader(f)
                # Extract text from all pages (limit to first 10 pages for performance)
                max_pages = min(len(pdf_reader.pages), 10)
                for page_num in range(max_pages):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            return ' '.join(text_parts) if text_parts else None
        except Exception as e:
            logging.debug(f"PDF extraction error for {file_path}: {e}")
            return None
    
    def _extract_docx_text(self, file_path):
        """Extract text from Word document using python-docx.
        
        Args:
            file_path: Path object for the Word document
            
        Returns:
            str: Extracted text content, or None on error
        """
        if not PYTHON_DOCX_AVAILABLE:
            return None
            
        try:
            doc = Document(str(file_path))
            text_parts = []
            # Extract text from all paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return ' '.join(text_parts) if text_parts else None
        except Exception as e:
            logging.debug(f"Word document extraction error for {file_path}: {e}")
            return None
    
    def extract_features(self, file_path):
        """Extract features from a file for classification including content analysis
        
        Args:
            file_path: Path to the file to extract features from
            
        Returns:
            str: Space-joined features extracted from the file
        """
        file_path = Path(file_path)
        
        # Create cache key using file path and content hash (if file exists)
        # This eliminates race conditions from mtime-based caching
        try:
            if file_path.exists():
                file_hash = self._compute_file_hash(file_path)
                cache_key = f"{file_path}:{file_hash}" if file_hash else None
                
                # Check cache first
                if cache_key is not None:
                    cached_features = self.feature_cache.get(cache_key)
                    if cached_features is not None:
                        logging.debug(f"Cache hit for {file_path}")
                        return cached_features
            else:
                cache_key = None
        except Exception as e:
            logging.debug(f"Error checking cache for {file_path}: {e}")
            cache_key = None
        
        features = []
        
        # Add filename as a feature
        features.append(file_path.name)
        
        # Add stem (filename without extension)
        features.append(file_path.stem)
        
        # Extract words from filename
        words = re.findall(r'[a-zA-Z]+', file_path.stem)
        features.extend(words)
        
        # Check if file exists before attempting to extract content
        if not file_path.exists():
            logging.warning(f"File does not exist: {file_path}")
            return ' '.join(features)
            
        # Try to extract content from various file types
        content_extracted = False
        file_extension = file_path.suffix.lower()
        
        # Try modern extraction methods based on file type
        try:
            # Check content cache first to avoid slow processing
            cached_content = self.content_cache.get_content(file_path)
            if cached_content is not None:
                logging.debug(f"Content cache hit for {file_path}")
                features.append(cached_content)
                content_extracted = True
            else:
                # SAFETY CHECK: Check file size before processing
                file_size_bytes = file_path.stat().st_size
                file_size_mb = file_size_bytes / (1024 * 1024)  # Convert to MB
                
                # EMPTY FILE CHECK: Warn about empty files (0 bytes)
                if file_size_bytes == 0:
                    logging.warning(f"Empty file detected (0 bytes): {file_path}")
                    warning_msg = "Empty file (0 bytes) - categorization based on filename only"
                    self.extraction_warnings[str(file_path)] = warning_msg
                    failure_info = {
                        'file': str(file_path),
                        'error': 'File is empty (0 bytes)',
                        'timestamp': datetime.now().isoformat(),
                        'reason': 'empty_file'
                    }
                    self.extraction_failures.append(failure_info)
                elif file_size_mb > self.max_file_size_mb:
                    # File too large - skip content extraction
                    logging.info(f"Skipping extraction for {file_path}: file size ({file_size_mb:.1f}MB) exceeds limit ({self.max_file_size_mb}MB)")
                    warning_msg = f"File too large ({file_size_mb:.1f}MB) - categorization based on filename only"
                    self.extraction_warnings[str(file_path)] = warning_msg
                    failure_info = {
                        'file': str(file_path),
                        'error': f'File size {file_size_mb:.1f}MB exceeds limit',
                        'timestamp': datetime.now().isoformat(),
                        'reason': 'size_limit'
                    }
                    self.extraction_failures.append(failure_info)
                else:
                    # File size OK - process with timeout protection
                    extracted_text = None
                    
                    try:
                        # PDF files
                        if file_extension == '.pdf':
                            extracted_text = _run_with_timeout(
                                self._extract_pdf_text,
                                args=(file_path,),
                                timeout_seconds=self.extraction_timeout_seconds
                            )
                        # Word documents
                        elif file_extension in ['.docx', '.doc']:
                            extracted_text = _run_with_timeout(
                                self._extract_docx_text,
                                args=(file_path,),
                                timeout_seconds=self.extraction_timeout_seconds
                            )
                        
                        if extracted_text:
                            # Limit content size to avoid memory issues
                            text_content = extracted_text[:5000]
                            # Cache the extracted content for future use
                            self.content_cache.set_content(file_path, text_content)
                            features.append(text_content)
                            content_extracted = True
                            logging.debug(f"Successfully extracted and cached content from {file_path}")
                    except TimeoutError as e:
                        logging.warning(f"Content extraction timeout for {file_path}: {e}")
                        warning_msg = f"Processing timeout ({self.extraction_timeout_seconds}s) - categorization based on filename only"
                        self.extraction_warnings[str(file_path)] = warning_msg
                        failure_info = {
                            'file': str(file_path),
                            'error': str(e),
                            'timestamp': datetime.now().isoformat(),
                            'reason': 'timeout'
                        }
                        self.extraction_failures.append(failure_info)
        except Exception as e:
            logging.warning(f"Content extraction failed for {file_path}: {e}")
            # Track this failure for user notification
            failure_info = {
                'file': str(file_path),
                'error': str(e),
                'timestamp': datetime.now().isoformat(),
                'reason': 'extraction_error'
            }
            self.extraction_failures.append(failure_info)
            self.extraction_warnings[str(file_path)] = "Content extraction failed - categorization based on filename only"
        
        # Fall back to basic text extraction if content wasn't extracted
        if not content_extracted and self._is_text_file(file_path):
            # BINARY FILE CHECK: Detect binary files before attempting text extraction
            if self._is_binary_file(file_path):
                logging.info(f"Binary file detected, skipping content extraction: {file_path}")
                warning_msg = "Binary file detected - categorization based on filename only"
                self.extraction_warnings[str(file_path)] = warning_msg
                failure_info = {
                    'file': str(file_path),
                    'error': 'Binary file detected',
                    'timestamp': datetime.now().isoformat(),
                    'reason': 'binary_file'
                }
                self.extraction_failures.append(failure_info)
            else:
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read(2000)  # Read first 2000 chars
                        features.append(content)
                        logging.debug(f"Successfully extracted content with basic method from {file_path}")
                except Exception as e:
                    logging.debug(f"Could not extract content from {file_path}: {e}")
        
        # Join features and cache the result
        result = ' '.join(features)
        
        # Store in cache if we have a valid cache key
        if cache_key is not None:
            self.feature_cache.put(cache_key, result)
            logging.debug(f"Cached features for {file_path}")
        
        return result
    
    def _is_binary_file(self, file_path):
        """Check if file is binary by reading first chunk and checking for null bytes.
        
        This is a heuristic approach - binary files typically contain null bytes
        while text files don't. We read the first 8KB to detect binary content.
        
        Args:
            file_path: Path to the file to check
            
        Returns:
            bool: True if file appears to be binary, False if it appears to be text
        """
        try:
            # Read first 8KB of file
            with open(file_path, 'rb') as f:
                chunk = f.read(8192)
            
            # Empty file is not binary
            if not chunk:
                return False
            
            # Check for null bytes (common in binary files, rare in text)
            if b'\x00' in chunk:
                return True
            
            # Check percentage of non-printable characters
            # Text files should have mostly printable characters
            non_printable_count = sum(1 for byte in chunk if byte < 32 and byte not in (9, 10, 13))
            non_printable_ratio = non_printable_count / len(chunk)
            
            # If more than 30% non-printable, likely binary
            return non_printable_ratio > 0.3
            
        except Exception as e:
            logging.debug(f"Error checking if file is binary {file_path}: {e}")
            # On error, assume text to allow processing to continue
            return False
    
    def _is_text_file(self, file_path):
        """Check if file is likely a text file based on extension
        
        Args:
            file_path: Path to the file to check
            
        Returns:
            bool: True if the file has a known text extension, False otherwise
        """
        text_extensions = [
            # Documentation formats
            '.txt', '.md', '.rst', '.rtf', '.tex', '.adoc', '.wiki',
            # Data formats
            '.csv', '.tsv', '.json', '.xml', '.yaml', '.yml', '.toml', '.ini', '.cfg', '.conf',
            # Web formats
            '.html', '.htm', '.css', '.scss', '.sass', '.less', '.js', '.jsx', '.ts', '.tsx',
            # Programming languages
            '.py', '.java', '.c', '.cpp', '.h', '.hpp', '.cs', '.php', '.rb', '.pl', '.swift',
            '.go', '.rs', '.kt', '.scala', '.sh', '.bash', '.ps1', '.bat', '.sql',
            # Log and config files
            '.log', '.properties', '.env'
        ]
        return file_path.suffix.lower() in text_extensions
    
    def train(self, file_paths, categories):
        """Train the classifier on file examples
        
        Args:
            file_paths: List of file paths
            categories: List of corresponding categories
        """
        if len(file_paths) != len(categories):
            raise ValueError("Number of files and categories must match")
            
        # Extract features for each file
        features = [self.extract_features(f) for f in file_paths]
        
        # Store unique categories
        self.categories = list(set(categories))
        
        # Train the model (different API for SentenceTransformer vs sklearn)
        if self.classifier_type == 'sentence_transformer':
            self.model.train(features, categories)
        else:
            self.model.fit(features, categories)
        
        self.trained = True
        logging.info(f"Trained AI classifier on {len(file_paths)} files with {len(self.categories)} categories")
    
    def predict(self, file_path):
        """Predict category for a file with confidence score
        
        Args:
            file_path: Path to the file to categorize
            
        Returns:
            dict: Prediction result containing category and confidence score,
                  or None if model is not trained
        """
        if not self.trained:
            raise ValueError("AI classifier must be trained before prediction. Call train() first.")
        
        # Check if file exists
        if not Path(file_path).exists():
            logging.warning(f"File does not exist: {file_path}")
            return {
                'category': 'Unknown',
                'confidence': 0.0,
                'error': 'File not found'
            }
            
        try:
            # Extract features
            features = self.extract_features(file_path)
            
            # Handle different classifier types
            if self.classifier_type == 'sentence_transformer':
                # SentenceTransformer returns (category, confidence, probabilities_dict)
                category, confidence, probabilities_dict = self.model.predict(features)
                
                # Get top 3 alternatives
                alternatives = []
                sorted_probs = sorted(probabilities_dict.items(), key=lambda x: x[1], reverse=True)
                for cat, prob in sorted_probs:
                    if cat != category and len(alternatives) < 3:
                        alternatives.append({
                            'category': cat,
                            'confidence': round(float(prob), 3)
                        })
            else:
                # sklearn Pipeline API
                category = self.model.predict([features])[0]
                probabilities = self.model.predict_proba([features])[0]
                
                # Find the confidence score (probability) for the predicted category
                category_index = list(self.model.classes_).index(category)
                confidence = probabilities[category_index]
                
                # Get top 3 alternative categories with their confidence scores
                alternatives = []
                sorted_indices = probabilities.argsort()[::-1]  # Sort indices by probability (descending)
                
                for i in sorted_indices[:3]:  # Get top 3
                    if self.model.classes_[i] != category:  # Skip the main prediction
                        alternatives.append({
                            'category': self.model.classes_[i],
                            'confidence': round(float(probabilities[i]), 3)
                        })
            
            result = {
                'category': category,
                'confidence': round(float(confidence), 3),
                'alternatives': alternatives
            }
            
            # Include extraction warning if content extraction failed for this file
            file_path_str = str(file_path)
            if file_path_str in self.extraction_warnings:
                result['extraction_warning'] = self.extraction_warnings[file_path_str]
                result['degraded_accuracy'] = True
                logging.info(f"WARNING: {file_path} categorized with degraded accuracy due to content extraction failure")
            
            return result
        except Exception as e:
            logging.error(f"Error predicting category for {file_path}: {e}")
            return {
                'category': 'Unknown',
                'confidence': 0.0,
                'error': str(e)
            }
    
    # ------------------------------------------------------------------
    # Bootstrap + progressive training
    # ------------------------------------------------------------------

    def _bootstrap_train(self):
        """Train on the built-in keyword corpus — called at __init__ time.

        Uses ``_BOOTSTRAP_CORPUS`` (module-level constant) so the model is
        always usable from the moment it is constructed, with no real files
        and no user intervention required.
        """
        texts      = [text for text, _ in _BOOTSTRAP_CORPUS]
        categories = [cat  for _, cat  in _BOOTSTRAP_CORPUS]
        try:
            if self.classifier_type == 'sentence_transformer':
                self.model.train(texts, categories)
            else:
                self.model.fit(texts, categories)
            self.categories = list(set(categories))
            self.trained = True
            logging.debug(
                "AIFileClassifier bootstrapped on built-in corpus (%d samples, %d categories)",
                len(texts), len(self.categories)
            )
        except Exception as e:
            # Bootstrap failure must never crash the app — log and move on.
            logging.warning("Bootstrap training failed (%s); model is untrained until real data arrives.", e)

    def safe_predict(self, file_path, confidence_threshold: float = 0.6) -> dict | None:
        """Predict category for *file_path*, gated by a confidence threshold.

        Unlike :meth:`predict`, this method never raises.  It returns
        ``None`` whenever the model is not trained or the top confidence
        score is below *confidence_threshold* — so callers can fall
        through to the next strategy without extra try/except.

        Args:
            file_path: Path to the file to classify.
            confidence_threshold: Minimum confidence score to accept the
                prediction (0–1). Default 0.6 — conservative enough to
                avoid false positives on the bootstrap prototype.

        Returns:
            dict with ``category`` and ``confidence`` keys, or ``None``.
        """
        if not self.trained:
            return None
        try:
            result = self.predict(file_path)
            if result and result.get('confidence', 0.0) >= confidence_threshold:
                return result
            return None
        except Exception as e:
            logging.debug("safe_predict: skipping %s — %s", file_path, e)
            return None

    def auto_train_if_needed(self, base_dir=None, model_path: str | None = None) -> bool:
        """Progressive learning: re-train from already-sorted files in *base_dir*.

        Scans *base_dir* recursively (using the same folder-structure label
        convention as :meth:`train_from_directory`) and rebuilds the model
        if at least 3 samples per category are found.  Falls back silently
        when there is not enough data — the bootstrap model keeps working.

        Args:
            base_dir: Root of the sorted output tree (e.g. ``~/Sortify``). 
                      Pass ``None`` to skip progressive learning and keep
                      the bootstrap model.
            model_path: If provided, save the freshly trained model here
                        so the next run can skip re-training.

        Returns:
            ``True`` if a progressive model was trained, ``False`` otherwise.
        """
        if base_dir is None:
            logging.debug("auto_train_if_needed: no base_dir supplied, keeping bootstrap model.")
            return False

        base_dir = Path(base_dir)
        if not base_dir.exists():
            logging.debug("auto_train_if_needed: base_dir does not exist: %s", base_dir)
            return False

        try:
            result = self.train_from_directory(base_dir, min_samples_per_category=3)
            if result.get('success', False):
                logging.info(
                    "Progressive training complete: %d files, %.1f%% accuracy.",
                    result.get('total_files', 0),
                    result.get('accuracy', 0.0) * 100,
                )
                if model_path:
                    self.save_model(model_path)
                return True
            else:
                logging.debug(
                    "auto_train_if_needed: not enough sorted data yet — %s",
                    result.get('error', 'unknown reason')
                )
                return False
        except Exception as e:
            logging.warning("auto_train_if_needed failed (%s); bootstrap model retained.", e)
            return False

    def save_model(self, model_path):
        """Save the trained model to disk"""
        if not self.trained:
            logging.warning("Cannot save untrained model")
            return False
            
        try:
            # Create directory if it doesn't exist (only if model_path has a directory component)
            model_dir = os.path.dirname(model_path)
            if model_dir:  # Only create dir if path has a directory component
                os.makedirs(model_dir, exist_ok=True)
            
            # Save model and categories
            with open(model_path, 'wb') as f:
                pickle.dump((self.model, self.categories), f)
                
            logging.info(f"Saved AI classifier model to {model_path}")
            return True
        except Exception as e:
            logging.error(f"Error saving model: {e}")
            return False
    
    def load_model(self, model_path):
        """Load a trained model from disk"""
        try:
            with open(model_path, 'rb') as f:
                self.model, self.categories = pickle.load(f)
                self.trained = True
                
            logging.info(f"Loaded AI classifier model from {model_path}")
            return True
        except Exception as e:
            logging.error(f"Error loading model: {e}")
            return False

    def evaluate(self, test_file_paths, test_categories):
        """Evaluate model performance on test data with detailed metrics
        
        Args:
            test_file_paths: List of file paths for testing
            test_categories: List of corresponding categories
            
        Returns:
            dict: Dictionary containing various performance metrics
                 (accuracy, precision, recall, f1_score)
        """
        if not self.trained:
            logging.warning("Cannot evaluate untrained model")
            return {'accuracy': 0.0}
            
        try:
            # Import metrics functions
            from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score, classification_report
            
            # Extract features for test files
            features = [self.extract_features(f) for f in test_file_paths]
            
            # Get predictions
            predictions = self.model.predict(features)
            
            # Calculate metrics
            accuracy = accuracy_score(test_categories, predictions)
            
            # For multi-class classification, we use weighted averaging
            precision = precision_score(test_categories, predictions, average='weighted', zero_division=0)
            recall = recall_score(test_categories, predictions, average='weighted', zero_division=0)
            f1 = f1_score(test_categories, predictions, average='weighted', zero_division=0)
            
            # Generate detailed report
            report = classification_report(test_categories, predictions, output_dict=True, zero_division=0)
            
            # Log results
            logging.info(f"AI classifier evaluation results:")
            logging.info(f"  Accuracy:  {accuracy:.3f}")
            logging.info(f"  Precision: {precision:.3f}")
            logging.info(f"  Recall:    {recall:.3f}")
            logging.info(f"  F1 Score:  {f1:.3f}")
            
            # Return comprehensive metrics
            return {
                'accuracy': round(float(accuracy), 3),
                'precision': round(float(precision), 3),
                'recall': round(float(recall), 3),
                'f1_score': round(float(f1), 3),
                'report': report,
                'categories': list(set(test_categories)),
                'test_size': len(test_file_paths)
            }
        except Exception as e:
            logging.error(f"Error during model evaluation: {e}")
            return {'accuracy': 0.0, 'error': str(e)}

    def get_extraction_stats(self):
        """Get statistics about content extraction failures.
        
        Returns:
            dict: Statistics including total failures, recent failures, and failure rate
        """
        total_failures = len(self.extraction_failures)
        
        # Get recent failures (last 10)
        recent_failures = self.extraction_failures[-10:] if total_failures > 0 else []
        
        # Calculate failure rate if we have cache stats
        feature_cache_stats = self.feature_cache.get_stats()
        total_extractions = feature_cache_stats['hits'] + feature_cache_stats['misses']
        failure_rate = (total_failures / total_extractions * 100) if total_extractions > 0 else 0
        
        return {
            'total_failures': total_failures,
            'recent_failures': recent_failures,
            'failure_rate': round(failure_rate, 2),
            'files_with_warnings': len(self.extraction_warnings),
            'has_degraded_accuracy': total_failures > 0
        }
    
    def clear_extraction_warnings(self):
        """Clear all extraction failure tracking data."""
        self.extraction_failures.clear()
        self.extraction_warnings.clear()
        logging.info("Cleared extraction failure tracking data")

    def train_with_split(self, file_paths, categories, test_size=0.2):
        """Train model with automatic train/test split and return detailed evaluation metrics
        
        Args:
            file_paths: List of file paths for training
            categories: List of corresponding categories
            test_size: Fraction of data to use for testing (default: 0.2)
            
        Returns:
            dict: Dictionary containing model performance metrics
        """
        if len(file_paths) < 5:
            logging.warning("Too few samples for reliable training and evaluation")
            return {'accuracy': 0.0, 'error': 'Too few samples'}
            
        try:
            # Split data into training and testing sets
            X_train, X_test, y_train, y_test = train_test_split(
                file_paths, categories, test_size=test_size, random_state=42, stratify=categories if len(set(categories)) > 1 else None
            )
            
            # Log split information
            logging.info(f"Split data into {len(X_train)} training and {len(X_test)} testing samples")
            
            # Train the model
            self.train(X_train, y_train)
            
            # Evaluate and get detailed metrics
            metrics = self.evaluate(X_test, y_test)
            
            # Add training information to metrics
            metrics['training_samples'] = len(X_train)
            metrics['testing_samples'] = len(X_test)
            metrics['categories_count'] = len(set(categories))
            
            return metrics
        except Exception as e:
            logging.error(f"Error during train_with_split: {e}")
            return {'accuracy': 0.0, 'error': str(e)}
        
    def train_from_directory(self, directory_path, test_size=0.2, min_samples_per_category=3, recursive=True, label_depth=None):
        """Train the classifier using files from a directory structure
        
        By default, this will recursively scan subfolders under the root and use the
        relative folder path as the category label (e.g., "ai_images/chatgpt").
        
        Args:
            directory_path: Path to the root directory containing category subdirectories
            test_size: Fraction of data to use for testing (default: 0.2)
            min_samples_per_category: Minimum number of samples required per category
            recursive: Whether to scan nested subfolders (default: True)
            label_depth: Limit label to the first N path components relative to root (default: None = full path)
        
        Returns:
            dict: Dictionary containing model performance metrics or
                  error information if training failed
        """
        directory_path = Path(directory_path)
        if not directory_path.exists() or not directory_path.is_dir():
            logging.error(f"Training directory not found: {directory_path}")
            return {'error': 'Training directory not found', 'success': False}
            
        file_paths = []
        categories = []
        category_counts = {}
        
        # Track skipped files with detailed reasons
        skipped_files = {
            'broken_symlinks': [],
            'special_files': [],
            'access_errors': [],
            'other': []
        }
        total_scanned = 0
        
        if recursive:
            # Recursively walk all files and derive labels from relative parent path
            for file_path in directory_path.rglob('*'):
                total_scanned += 1
                
                # Validate file type and accessibility
                try:
                    # Check for broken symlinks (symlink that points to non-existent target)
                    if file_path.is_symlink() and not file_path.exists():
                        logging.debug(f"Skipping broken symlink: {file_path}")
                        skipped_files['broken_symlinks'].append(str(file_path))
                        continue
                    
                    # Check if it's a regular file (not directory, device, pipe, socket, etc.)
                    if not file_path.is_file():
                        # It's either a directory or special file
                        if not file_path.is_dir():
                            # It's a special file (device, pipe, socket, etc.)
                            logging.debug(f"Skipping special file: {file_path}")
                            skipped_files['special_files'].append(str(file_path))
                        # Directories are expected, don't log them as skipped
                        continue
                        
                except PermissionError as e:
                    logging.debug(f"Skipping file due to permission error: {file_path} - {e}")
                    skipped_files['access_errors'].append(str(file_path))
                    continue
                except Exception as e:
                    logging.debug(f"Skipping file due to unexpected error: {file_path} - {e}")
                    skipped_files['other'].append(str(file_path))
                    continue
                
                # File is valid, extract category
                try:
                    rel_parent = file_path.parent.relative_to(directory_path)
                except ValueError:
                    # Should not happen, but guard anyway
                    continue
                    
                # Ignore files that are directly under the root (no category)
                if rel_parent == Path('.') or str(rel_parent).strip() == '':
                    continue
                    
                parts = rel_parent.parts
                if label_depth is not None and isinstance(label_depth, int) and label_depth > 0:
                    parts = parts[:label_depth]
                category = '/'.join(parts)
                file_paths.append(str(file_path))
                categories.append(category)
                category_counts[category] = category_counts.get(category, 0) + 1
        else:
            # Backward-compatible non-recursive behavior (immediate subfolders only)
            for category_dir in directory_path.iterdir():
                if category_dir.is_dir():
                    category = category_dir.name
                    category_files = []
                    for file_path in category_dir.glob('*'):
                        total_scanned += 1
                        
                        # Validate file type and accessibility
                        try:
                            # Check for broken symlinks
                            if file_path.is_symlink() and not file_path.exists():
                                logging.debug(f"Skipping broken symlink: {file_path}")
                                skipped_files['broken_symlinks'].append(str(file_path))
                                continue
                            
                            # Check if it's a regular file
                            if not file_path.is_file():
                                if not file_path.is_dir():
                                    logging.debug(f"Skipping special file: {file_path}")
                                    skipped_files['special_files'].append(str(file_path))
                                continue
                                
                        except PermissionError as e:
                            logging.debug(f"Skipping file due to permission error: {file_path} - {e}")
                            skipped_files['access_errors'].append(str(file_path))
                            continue
                        except Exception as e:
                            logging.debug(f"Skipping file due to unexpected error: {file_path} - {e}")
                            skipped_files['other'].append(str(file_path))
                            continue
                        
                        file_paths.append(str(file_path))
                        categories.append(category)
                        category_files.append(str(file_path))
                    category_counts[category] = len(category_files)
        
        # Calculate total skipped files
        total_skipped = sum(len(files) for files in skipped_files.values())
        
        # Warn user if files were skipped
        if total_skipped > 0:
            skipped_breakdown = []
            if skipped_files['broken_symlinks']:
                skipped_breakdown.append(f"{len(skipped_files['broken_symlinks'])} broken symlinks")
            if skipped_files['special_files']:
                skipped_breakdown.append(f"{len(skipped_files['special_files'])} special files")
            if skipped_files['access_errors']:
                skipped_breakdown.append(f"{len(skipped_files['access_errors'])} permission errors")
            if skipped_files['other']:
                skipped_breakdown.append(f"{len(skipped_files['other'])} other errors")
            
            breakdown_str = ", ".join(skipped_breakdown)
            logging.warning(
                f"Training data incomplete: {total_skipped} file(s) skipped out of {total_scanned} scanned "
                f"({breakdown_str}). Model accuracy may be degraded. "
                f"Please review skipped files in training metrics."
            )
        
        if not file_paths:
            logging.warning(f"No training files found in {directory_path}")
            return {
                'error': 'No training files found',
                'success': False,
                'skipped_files': skipped_files,
                'total_skipped': total_skipped,
                'total_scanned': total_scanned
            }
        
        # Check if we have enough samples per category
        insufficient_categories = [cat for cat, count in category_counts.items() if count < min_samples_per_category]
        if insufficient_categories:
            logging.warning(f"Insufficient samples for categories: {', '.join(insufficient_categories)}")
            logging.warning(f"Each category needs at least {min_samples_per_category} samples for reliable training")
            return {
                'error': f"Insufficient samples for {len(insufficient_categories)} categories",
                'insufficient_categories': insufficient_categories,
                'category_counts': category_counts,
                'success': False,
                'skipped_files': skipped_files,
                'total_skipped': total_skipped,
                'total_scanned': total_scanned
            }
            
        # Train the model with the collected files
        logging.info(
            f"Training from directory with {len(file_paths)} files in {len(set(categories))} categories (recursive={recursive}, label_depth={label_depth})"
        )
        metrics = self.train_with_split(file_paths, categories, test_size)
        
        # Add category distribution information to metrics
        metrics['category_distribution'] = category_counts
        metrics['success'] = True
        
        # Add skipped file information to metrics
        metrics['skipped_files'] = skipped_files
        metrics['total_skipped'] = total_skipped
        metrics['total_scanned'] = total_scanned
        metrics['valid_files_processed'] = len(file_paths)
        
        return metrics
